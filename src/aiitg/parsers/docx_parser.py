"""docx parser: python-docx + raw OOXML → unified text model.

Reads body paragraphs, headers/footers, and tables, capturing per-run style
(font size, color, w:vanish) so detectors can spot hidden/invisible text.
Also exposes raw OOXML parts for the raw-node detector.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from aiitg.core.document import ParsedDocument, TextParagraph, TextRun
from aiitg.parsers.base import register_parser
from aiitg.parsers.ooxml_raw import OOXMLRawReader

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _normalize_color(val: str | None) -> str | None:
    """Normalize a color to RRGGBB (uppercase, 6 hex digits) or None."""
    if not val:
        return None
    v = val.strip().lstrip("#").upper()
    if re.fullmatch(r"[0-9A-F]{6}", v):
        return v
    if re.fullmatch(r"[0-9A-F]{3}", v):
        return "".join(c * 2 for c in v)
    return None


def _run_is_vanish(run_xml) -> bool:
    """True if the run carries w:vanish or w:webHidden (hidden text)."""
    rpr = run_xml.find(qn("w:rPr"))
    if rpr is None:
        return False
    return rpr.find(qn("w:vanish")) is not None or rpr.find(qn("w:webHidden")) is not None


def _run_color(run_xml) -> str | None:
    rpr = run_xml.find(qn("w:rPr"))
    if rpr is None:
        return None
    color_el = rpr.find(qn("w:color"))
    if color_el is None:
        return None
    return _normalize_color(color_el.get(qn("w:val")))


def _run_size_halfpoints(run_xml) -> float | None:
    """Font size in half-points (w:sz), or None."""
    rpr = run_xml.find(qn("w:rPr"))
    if rpr is None:
        return None
    sz = rpr.find(qn("w:sz"))
    if sz is None:
        return None
    val = sz.get(qn("w:val"))
    return float(val) / 2.0 if val else None


def _extract_run(paragraph_xml, run, run_index: int) -> TextRun:
    """Build a TextRun from a python-docx run + its underlying XML."""
    text = run.text or ""
    try:
        font = run.font
        size_pt = font.size.pt if font.size else None
        color_val = _normalize_color(font.color.rgb) if font.color and font.color.type is not None else None
    except Exception:  # noqa: BLE001 — style access can fail on weird docs
        size_pt, color_val = None, None

    raw = {
        "xml_index": run_index,
        "italic": bool(run.italic),
        "bold": bool(run.bold),
    }
    return TextRun(
        text=text,
        font_size=size_pt,
        color=color_val,
        is_hidden=_run_is_vanish(run._element),
        raw=raw,
    )


def _paragraph_from_xml(p_el, index: int | None, is_hdr_ftr: bool = False) -> TextParagraph:
    """Build a TextParagraph from the raw w:p element (keeps hidden text)."""
    runs: list[TextRun] = []
    for i, r_el in enumerate(p_el.findall(qn("w:r"))):
        # w:t text
        text = "".join(t.text or "" for t in r_el.findall(qn("w:t")))
        if not text and r_el.find(qn("w:delText")) is not None:
            # tracked-deleted text is still visible to a machine reader
            text = "".join(d.text or "" for d in r_el.findall(qn("w:delText")))
        if not text:
            continue
        runs.append(
            TextRun(
                text=text,
                font_size=_run_size_halfpoints(r_el),
                color=_run_color(r_el),
                is_hidden=_run_is_vanish(r_el),
                is_header_footer=is_hdr_ftr,
                raw={"xml_index": i},
            )
        )
    full_text = "".join(r.text for r in runs)
    return TextParagraph(text=full_text, index=index, runs=runs)


@register_parser("docx", ("docx",), sniff=None)
class DocxParser:
    """Parse .docx into :class:`ParsedDocument`."""

    def parse(self, path: Path) -> ParsedDocument:
        doc = DocxDocument(str(path))
        doc_model = ParsedDocument(kind="docx", source_path=str(path))

        # body paragraphs (with raw XML to catch hidden text)
        body = doc.element.body
        idx = 0
        for p_el in body.iter(qn("w:p")):
            para = _paragraph_from_xml(p_el, idx)
            if para.text:
                doc_model.paragraphs.append(para)
                idx += 1

        # headers/footers (python-docx sections)
        try:
            for section in doc.sections:
                for header in (section.header, section.footer):
                    if not header.is_linked_to_previous:
                        for p_el in header._element.iter(qn("w:p")):
                            para = _paragraph_from_xml(p_el, None, is_hdr_ftr=True)
                            if para.text:
                                doc_model.paragraphs.append(para)
        except Exception:  # noqa: BLE001 — header/footer access can be flaky
            doc_model.warnings.append("failed to read headers/footers")

        # tables (cells → paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text:
                            doc_model.paragraphs.append(
                                TextParagraph(text=p.text, index=None, runs=[])
                            )

        # metadata (core properties)
        props = doc.core_properties
        doc_model.metadata = {
            "author": props.author,
            "title": props.title,
            "subject": props.subject,
            "keywords": props.keywords,
            "comments": props.comments,
            "last_modified_by": props.last_modified_by,
        }

        # raw OOXML parts (for raw-node detector)
        try:
            with OOXMLRawReader(path) as reader:
                doc_model.ooxml_parts = {name: reader.part_bytes(name) for name in reader.part_names()}
        except Exception as exc:  # noqa: BLE001
            doc_model.warnings.append(f"failed to read raw OOXML parts: {exc}")

        return doc_model
